"""
markdown_to_docx.py

Utility functions for processing markdown text, converting it to
DOCX format, and uploading the resulting file to Slack.

Functions:
- markdown_to_docx: Converts markdown-formatted text to a DOCX file.
- save_report_to_docx: Converts a report in markdown format to a
    DOCX file.
- send_docx_to_slack: Sends a DOCX file to a specified Slack channel.

Attributes:
    None
"""

import re
import tempfile
import time

from pathlib import Path
from docx import Document
from docx.shared import Pt
from slack_sdk.errors import SlackApiError

from utils.logging_utils import log_message, log_error

def markdown_to_docx(
        text: str,
        output_path: str,
) -> None:
    """
    Converts markdown-formatted text to a DOCX file and saves it to a
    specified path. Supports main headings, sub-headings, bullet points,
    bold, and italic formatting.

    Args:
        text (str): The markdown-formatted text to convert.
        output_path (str): The path to save the resulting DOCX file.

    Returns:
        None
    """
    doc = Document()

    lines = text.split('\n')

    for line in lines:
        line = line.rstrip()

        # Handle headings (from # to ######)
        if line.startswith('#'):
            # Count the number of '#' symbols
            heading_level = len(line.split(' ')[0])
            heading_text = line[heading_level:].strip()

            # Add the heading based on the level
            if heading_level == 1:
                doc.add_heading(heading_text, level=1)
            elif heading_level == 2:
                doc.add_heading(heading_text, level=2)
            elif heading_level == 3:
                doc.add_heading(heading_text, level=3)
            elif heading_level == 4:
                doc.add_heading(heading_text, level=4)
            elif heading_level == 5:
                doc.add_heading(heading_text, level=5)
            elif heading_level == 6:
                doc.add_heading(heading_text, level=6)

        # Handle bullet points, considering indentation
        elif line.strip().startswith('-'):
            # Check if this is a horizontal rule
            # (line with only hyphens)
            if len(line.strip()) > 1 and len(
                line.strip()) == line.lstrip().count('-'):
                # Line with hyphens only (e.g., "----------")
                # Add empty paragraph for spacing
                doc.add_paragraph('', style='Normal')
                doc.add_paragraph('-' * 40)  # Add a horizontal line
                continue  # Skip further processing for horizontal line

            # Skip lines that start with a hyphen but no spaces
            # before it (for bullet points)
            if line.lstrip().startswith('-') and len(line) == len(
                line.lstrip()
            ):
                continue  # Skip if no indentation

            # Calculate indentation level based on leading spaces
            indent_level = (len(line) - len(line.lstrip())) // 2
            bullet_text = line.replace("-", "").strip()

            # Apply bullet styles based on indentation level
            if indent_level == 0:
                para = doc.add_paragraph(bullet_text, style='List Bullet')
            elif indent_level == 1:
                para = doc.add_paragraph(bullet_text, style='List Bullet 2')
            elif indent_level == 2:
                para = doc.add_paragraph(bullet_text, style='List Bullet 3')
            else:
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                para.paragraph_format.left_indent = Pt(20 * indent_level)

            # Adjust font size for nested levels for a clean look
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(10 if indent_level > 0 else 12)

        # Handle bold (**bold**) and italics (*italic* or _italic_)
        else:
            segments = re.split(r'(\*\*.*?\*\*|\*.*?\*|_.*?_)', line)
            para = doc.add_paragraph()

            for segment in segments:
                if segment.startswith('**') and segment.endswith('**'):
                    # Bold text
                    run = para.add_run(segment[2:-2])
                    run.bold = True
                elif (segment.startswith('*') and segment.endswith('*')) or (
                        segment.startswith('_') and segment.endswith('_')
                ):
                    # Italic text
                    run = para.add_run(segment[1:-1])
                    run.italic = True
                else:
                    # Regular text
                    para.add_run(segment)

    # Save the document
    doc.save(output_path)


def save_report_to_docx(report_text: str,) -> str:
    """
    Converts a report in markdown format to a DOCX file and saves it
    to a temporary location.

    Args:
        report_text (str): The report text in markdown format.

    Returns:
        str: The path to the temporary .docx file.
    """
    temp_file = tempfile.NamedTemporaryFile(
        delete=False, suffix=".docx"
    )
    temp_file.close()

    markdown_to_docx(report_text, temp_file.name)

    return temp_file.name


def send_docx_to_slack(
        client: object,
        channel_id: str,
        thread_ts: str,
        file_path: str,
) -> None:
    """
    Sends a DOCX file to a specified Slack channel and thread.

    Args:
        client (object): The Slack client instance.
        channel_id (str): The ID of the Slack channel.
        thread_ts (str): The timestamp of the thread message.
        file_path (str): The path to the DOCX file to upload.

    Returns:
        None
    
    Raises:
        SlackApiError: An error occurred when uploading the file.
    """
    attempt = 0
    max_retries = 3
    retry_delay = 5

    while attempt < max_retries:
        try:
            # Open the file and upload it to Slack
            with open(file_path, "rb") as file_content:
                # Extract the file name from the full path
                file_name = Path(file_path).name

                # Attempt the file upload
                response = client.files_upload_v2(
                    channel=channel_id,
                    thread_ts=thread_ts,
                    file=file_content,
                    filename=file_name,
                    title="Generated Report"
                )

                # Log success
                log_message(
                    f"File uploaded successfully: {response['file']['id']}",
                    "info"
                )
                return  # Exit the function if upload is successful

        except SlackApiError as e:
            attempt += 1
            log_error(e, f"Sending file to Slack: Attempt {attempt} failed.")
            if attempt < max_retries:
                log_message(f"Retrying in {retry_delay} seconds...", level="warning")
                time.sleep(retry_delay)
            else:
                log_message("Max retries reached. File upload failed.", level="error")
                raise
